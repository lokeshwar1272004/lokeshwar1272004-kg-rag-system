"""
ingestion/loader.py
Document ingestion: loads PDF, TXT, DOCX files and returns raw text.
Supports both file path and uploaded bytes input (for Streamlit).
"""

import io
from pathlib import Path
from typing import Union, List, Dict, Optional
from dataclasses import dataclass

from utils.logger import logger
from utils.text_utils import clean_text


@dataclass
class Document:
    """Represents a loaded document."""
    doc_id: str
    source: str
    content: str
    metadata: Dict


class DocumentLoader:
    """
    Multi-format document loader.
    Supports: PDF, TXT, DOCX, Markdown
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def load_file(self, file_path: Union[str, Path]) -> Optional[Document]:
        """Load a single file by path."""
        path = Path(file_path)

        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file type: {path.suffix}")
            return None

        logger.info(f"Loading file: {path.name}")

        try:
            if path.suffix.lower() == ".pdf":
                content = self._load_pdf(path)
            elif path.suffix.lower() == ".docx":
                content = self._load_docx(path)
            else:
                content = self._load_text(path)

            cleaned = clean_text(content)

            return Document(
                doc_id=path.stem,
                source=str(path),
                content=cleaned,
                metadata={
                    "filename": path.name,
                    "extension": path.suffix,
                    "size_chars": len(cleaned),
                    "source_type": "file",
                }
            )
        except Exception as e:
            logger.error(f"Failed to load {path.name}: {e}")
            return None

    def load_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str = "pdf"
    ) -> Optional[Document]:
        """Load from raw bytes (Streamlit file upload)."""
        try:
            if file_type == "pdf" or filename.endswith(".pdf"):
                content = self._load_pdf_bytes(file_bytes)
            elif filename.endswith(".docx"):
                content = self._load_docx_bytes(file_bytes)
            else:
                content = file_bytes.decode("utf-8", errors="replace")

            cleaned = clean_text(content)

            return Document(
                doc_id=Path(filename).stem,
                source=filename,
                content=cleaned,
                metadata={
                    "filename": filename,
                    "size_chars": len(cleaned),
                    "source_type": "upload",
                }
            )
        except Exception as e:
            logger.error(f"Failed to load bytes for {filename}: {e}")
            return None

    def load_text_directly(self, text: str, source_name: str = "manual_input") -> Document:
        """Accept raw text directly (for demo/testing)."""
        cleaned = clean_text(text)
        return Document(
            doc_id=source_name,
            source=source_name,
            content=cleaned,
            metadata={
                "filename": source_name,
                "size_chars": len(cleaned),
                "source_type": "text_input",
            }
        )

    def load_directory(self, dir_path: Union[str, Path]) -> List[Document]:
        """Load all supported files from a directory."""
        path = Path(dir_path)
        docs = []

        if not path.is_dir():
            logger.error(f"Not a directory: {dir_path}")
            return docs

        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                doc = self.load_file(file_path)
                if doc:
                    docs.append(doc)

        logger.info(f"Loaded {len(docs)} documents from {dir_path}")
        return docs

    # ── Internal loaders ──────────────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> str:
        """Load PDF from file path."""
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    def _load_pdf_bytes(self, data: bytes) -> str:
        """Load PDF from bytes."""
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    def _load_docx(self, path: Path) -> str:
        """Load DOCX from file path."""
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    def _load_docx_bytes(self, data: bytes) -> str:
        """Load DOCX from bytes."""
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    def _load_text(self, path: Path) -> str:
        """Load plain text."""
        return path.read_text(encoding="utf-8", errors="replace")


# Singleton
document_loader = DocumentLoader()

__all__ = ["document_loader", "DocumentLoader", "Document"]
